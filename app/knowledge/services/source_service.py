import asyncio
from datetime import datetime
import hashlib
import logging
import mimetypes
from pathlib import Path
import shutil
from typing import Any, BinaryIO, ClassVar
import uuid

import aiofiles
from fastapi import UploadFile
from langchain.text_splitter import RecursiveCharacterTextSplitter
import openai

from app.core.settings import settings
from app.core.text_processing import TextProcessor
from app.knowledge.infrastructure.vector_store_client import VectorStoreClient
from app.knowledge.models.chunk import DocumentChunk
from app.knowledge.models.embedding import EmbeddingMetadata
from app.knowledge.models.source import (
    DocumentStatus,
    DocumentType,
    SourceDocument,
    SourceDocumentSummary,
)
from app.knowledge.services.embedding_service import EmbeddingService

logger = logging.getLogger(__name__)


class SourceServiceError(Exception):
    """Base exception for source service operations."""

    pass


class SourceService:
    # Define supported document formats and MIME types
    SUPPORTED_DOCUMENT_FORMATS: ClassVar = {
        "application/pdf": DocumentType.PDF,
        "text/plain": DocumentType.TXT,
        "text/markdown": DocumentType.MARKDOWN,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DocumentType.DOCX,
        "text/html": DocumentType.HTML,
        "application/json": DocumentType.JSON,
        "text/csv": DocumentType.CSV,
    }

    # Define supported audio formats
    SUPPORTED_AUDIO_FORMATS: ClassVar = {
        "audio/mpeg",  # mp3
        "audio/wav",  # wav
        "audio/x-wav",  # wav alternative
        "audio/mp4",  # m4a
        "audio/aac",  # aac
        "audio/ogg",  # ogg
        "audio/flac",  # flac
        "audio/webm",  # webm
    }

    AUDIO_EXTENSIONS: ClassVar = {
        ".mp3",
        ".wav",
        ".m4a",
        ".aac",
        ".ogg",
        ".flac",
        ".webm",
    }

    def __init__(
        self,
        text_processor: TextProcessor,
        embedding_service: EmbeddingService,
        vector_store_client: VectorStoreClient,
        upload_dir: str | None = None,
        openai_api_key: str | None = None,
    ):
        """
        Initialize the source service.

        Args:
            text_processor: Service for text processing
            embedding_service: Service for generating embeddings
            vector_store_client: Client for vector database operations
            upload_dir: Directory for storing uploaded files
            openai_api_key: OpenAI API key for transcription services

        """
        self.text_processor = text_processor
        self.embedding_service = embedding_service
        self.vector_store = vector_store_client

        # Use settings from config if not provided
        self.upload_dir = upload_dir or settings.upload_dir
        self.ensure_upload_dir_exists()

        # Configure OpenAI client for audio transcription if key provided
        self.openai_client = None
        if openai_api_key:
            openai.api_key = openai_api_key
            self.openai_client = openai

        # Text splitter configuration from settings
        self.chunk_size = settings.knowledge_config.document_processing.chunk_size
        self.chunk_overlap = settings.knowledge_config.document_processing.chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size, chunk_overlap=self.chunk_overlap
        )

    def ensure_upload_dir_exists(self):
        """Ensure that the upload directory exists."""
        Path(self.upload_dir).mkdir(parents=True, exist_ok=True)
        logger.info(f"Ensuring upload directory exists: {self.upload_dir}")

    def _determine_document_type(self, file_path: str, content_type: str) -> DocumentType:
        """
        Determine document type from MIME type and file extension.

        Args:
            file_path: Path to the file
            content_type: MIME type of the file

        Returns:
            Document type enumeration value

        Raises:
            SourceServiceError: If file type is not supported

        """
        # Check if it's an audio file
        if self._is_audio_file(file_path, content_type):
            return DocumentType.AUDIO

        # Check MIME type against supported document formats
        if content_type in self.SUPPORTED_DOCUMENT_FORMATS:
            return self.SUPPORTED_DOCUMENT_FORMATS[content_type]

        # Fallback to file extension
        file_extension = Path(file_path).suffix.lower()
        extension_map = {
            ".pdf": DocumentType.PDF,
            ".txt": DocumentType.TXT,
            ".md": DocumentType.MARKDOWN,
            ".docx": DocumentType.DOCX,
            ".html": DocumentType.HTML,
            ".json": DocumentType.JSON,
            ".csv": DocumentType.CSV,
        }

        if file_extension in extension_map:
            return extension_map[file_extension]

        raise SourceServiceError(f"Unsupported file type: {content_type}")

    def _is_audio_file(self, file_path: str, content_type: str) -> bool:
        """
        Check if the file is an audio file based on MIME type and extension.

        Args:
            file_path: Path to the file
            content_type: MIME type of the file

        Returns:
            True if the file is an audio file

        """
        # Check MIME type
        if content_type in self.SUPPORTED_AUDIO_FORMATS:
            return True

        # Check file extension as fallback
        file_extension = Path(file_path).suffix.lower()
        return file_extension in self.AUDIO_EXTENSIONS

    async def _transcribe_audio_file(self, file_path: str) -> str:
        """
        Transcribe audio file to text using OpenAI Whisper API.

        Args:
            file_path: Path to the audio file

        Returns:
            Transcribed text content

        Raises:
            SourceServiceError: If transcription fails

        """
        if not self.openai_client:
            raise SourceServiceError("OpenAI client not configured. Please provide openai_api_key.")

        try:
            logger.info(f"Starting transcription for audio file: {file_path}")

            # Check file size (OpenAI has a 25MB limit)
            file_size = Path(file_path).stat().st_size
            max_size = 25 * 1024 * 1024  # 25MB

            if file_size > max_size:
                raise SourceServiceError(f"Audio file exceeds 25MB limit ({file_size / 1024 / 1024:.2f}MB)")

            # Transcribe the audio
            with file_path.open("rb") as audio_file:
                transcript = await asyncio.to_thread(
                    self.openai_client.audio.transcriptions.create,
                    model="whisper-1",
                    file=audio_file,
                )

            # Extract text from response
            text = transcript.text if hasattr(transcript, "text") else transcript.get("text", "")

            if not text:
                raise SourceServiceError("Audio transcription returned empty text")

            logger.info(f"Transcription successful: {len(text)} characters")
            return text

        except Exception as e:
            logger.error(f"Transcription failed: {e!s}")
            raise SourceServiceError(f"Audio transcription failed: {e!s}") from e

    async def _extract_text_from_file(self, file_path: str, file_type: DocumentType) -> str:
        """
        Extract text content from document file.

        Args:
            file_path: Path to the document file
            file_type: Type of the document

        Returns:
            Extracted text content

        Raises:
            SourceServiceError: If text extraction fails

        """
        try:
            logger.info(f"Extracting text from {file_type} file: {file_path}")

            # Handle different file types
            if file_type == DocumentType.PDF:
                return await self._extract_text_from_pdf(file_path)
            if file_type in (DocumentType.TXT, DocumentType.MARKDOWN):
                return await self._extract_text_from_txt(file_path)
            if file_type == DocumentType.DOCX:
                return await self._extract_text_from_docx(file_path)
            if file_type in (DocumentType.HTML, DocumentType.JSON, DocumentType.CSV):
                return await self._extract_text_from_txt(file_path)
            if file_type == DocumentType.AUDIO:
                return await self._transcribe_audio_file(file_path)
            raise SourceServiceError(f"Unsupported file type for text extraction: {file_type}")

        except Exception as e:
            logger.error(f"Text extraction failed: {e!s}")
            raise SourceServiceError(f"Text extraction failed: {e!s}") from e

    async def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            from pypdf import PdfReader

            text = ""
            pdf = await asyncio.to_thread(PdfReader, file_path)

            for page in pdf.pages:
                page_text = await asyncio.to_thread(page.extract_text)
                if page_text:
                    text += page_text + "\n\n"

            if not text.strip():
                logger.warning(f"PDF extraction returned empty text: {file_path}")

            return text
        except Exception as e:
            logger.error(f"PDF extraction failed: {e!s}")
            raise SourceServiceError(f"PDF extraction failed: {e!s}") from e

    async def _extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from plain text file."""
        try:
            async with aiofiles.open(file_path, encoding="utf-8", errors="replace") as f:
                return await f.read()
        except Exception as e:
            logger.error(f"Text file reading failed: {e!s}")
            raise SourceServiceError(f"Text file reading failed: {e!s}") from e

    async def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            import docx

            doc = await asyncio.to_thread(docx.Document, file_path)
            text = ""

            for para in doc.paragraphs:
                text += para.text + "\n"

            return text
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e!s}")
            raise SourceServiceError(f"DOCX extraction failed: {e!s}") from e

    async def _compute_content_hash(self, file: BinaryIO) -> str:
        """Compute SHA-256 hash of file content for deduplication."""
        try:
            file_hash = hashlib.sha256()
            # Get current position to restore later
            pos = file.tell()
            # Go back to beginning of file
            file.seek(0)

            # Read and update hash in chunks to handle large files
            for byte_block in iter(lambda: file.read(4096), b""):
                file_hash.update(byte_block)

            # Restore original position
            file.seek(pos)

            return file_hash.hexdigest()
        except Exception as e:
            logger.error(f"Hash computation failed: {e!s}")
            raise SourceServiceError(f"Hash computation failed: {e!s}") from e

    async def _save_uploaded_file(
        self, upload_file: UploadFile, dest_dir: str, filename: str | None = None
    ) -> tuple[str, int]:
        """
        Save an uploaded file to disk.

        Args:
            upload_file: File uploaded via FastAPI
            dest_dir: Destination directory
            filename: Optional custom filename

        Returns:
            Tuple of (file path, file size in bytes)

        Raises:
            SourceServiceError: If file save fails

        """
        try:
            # Ensure destination directory exists
            Path(dest_dir).mkdir(parents=True, exist_ok=True)

            # Generate unique filename if not provided
            if not filename:
                ext = Path(upload_file.filename).suffix
                filename = f"{uuid.uuid4()}{ext}"

            file_path = Path(dest_dir) / filename

            # Save file
            async with aiofiles.open(file_path, "wb") as f:
                # Get current position
                pos = upload_file.file.tell()
                # Go back to beginning
                upload_file.file.seek(0)

                # Copy in chunks
                while chunk := upload_file.file.read(1024 * 1024):  # 1MB chunks
                    await f.write(chunk)

                # Restore position
                upload_file.file.seek(pos)

            # Get file size
            file_size = Path(file_path).stat().st_size

            logger.info(f"Saved uploaded file: {file_path} ({file_size} bytes)")
            return file_path, file_size

        except Exception as e:
            logger.error(f"File save failed: {e!s}")
            raise SourceServiceError(f"Failed to save uploaded file: {e!s}") from e

    async def _chunk_text(self, text: str, metadata: dict[str, Any]) -> list[dict[str, Any]]:
        """
        Split text into chunks for embedding.

        Args:
            text: Source text to split
            metadata: Metadata to include with each chunk

        Returns:
            List of chunk objects with text and metadata

        """
        try:
            # Split text into chunks
            chunks = self.text_splitter.create_documents([text])

            # Format chunks with metadata
            result = []
            for i, chunk in enumerate(chunks):
                chunk_text = chunk.page_content
                chunk_metadata = {**metadata, "chunk_index": i}

                if not chunk_text.strip():
                    logger.warning(f"Skipping empty chunk {i}")
                    continue

                result.append({"text": chunk_text, "metadata": chunk_metadata})

            logger.info(f"Text split into {len(result)} chunks")
            return result

        except Exception as e:
            logger.error(f"Text chunking failed: {e!s}")
            raise SourceServiceError(f"Text chunking failed: {e!s}") from e

    async def _store_chunks(
        self, source_id: str, chunks: list[dict[str, Any]]
    ) -> tuple[list[DocumentChunk], list[EmbeddingMetadata]]:
        """
        Generate embeddings and store chunks in vector database.

        Args:
            source_id: ID of the source document
            chunks: List of text chunks with metadata

        Returns:
            Tuple of (document chunks, embedding metadata)

        """
        try:
            # Prepare texts for embedding
            texts = [chunk["text"] for chunk in chunks]
            metadatas = [chunk["metadata"] for chunk in chunks]

            # Generate embeddings
            embeddings = await self.embedding_service.generate_embeddings(texts)

            if len(embeddings) != len(texts):
                raise SourceServiceError(f"Embedding count mismatch: expected {len(texts)}, got {len(embeddings)}")

            # Add to vector store
            chunk_ids = [str(uuid.uuid4()) for _ in chunks]

            logger.info(f"Adding {len(chunk_ids)} vectors to collection")

            # Add to vector database
            await self.vector_store.get_or_create_collection(settings.knowledge_config.vector_db.documents_collection)

            await self.vector_store.add_texts(
                collection_name=settings.knowledge_config.vector_db.documents_collection,
                texts=texts,
                embeddings=embeddings,
                metadatas=metadatas,
                ids=chunk_ids,
            )

            # Create DocumentChunk objects
            document_chunks = []
            embedding_metadata = []

            for i, (chunk_id, chunk) in enumerate(zip(chunk_ids, chunks, strict=False)):
                # Create chunk
                doc_chunk = DocumentChunk(
                    id=chunk_id,
                    source_document_id=source_id,
                    content=chunk["text"],
                    chunk_index=i,
                    metadata=chunk["metadata"],
                    vector_id=chunk_id,
                )
                document_chunks.append(doc_chunk)

                # Create embedding metadata
                embed_meta = EmbeddingMetadata(
                    source_document_id=source_id,
                    chunk_id=chunk_id,
                    vector_id=chunk_id,
                    model_name=self.embedding_service.model_name,
                    dimensions=self.embedding_service.dimension,
                )
                embedding_metadata.append(embed_meta)

            logger.info(f"Created {len(document_chunks)} chunk objects and metadata entries")
            return document_chunks, embedding_metadata

        except Exception as e:
            logger.error(f"Chunk storage failed: {e!s}")
            raise SourceServiceError(f"Chunk storage failed: {e!s}") from e

    async def _update_source_document_status(
        self, doc: SourceDocument, status: DocumentStatus, error: str | None = None
    ) -> SourceDocument:
        """
        Update the status of a source document.

        Args:
            doc: Source document to update
            status: New status
            error: Optional error message

        Returns:
            Updated source document

        """
        doc.status = status
        doc.updated_at = datetime.utcnow()

        if status == DocumentStatus.COMPLETED:
            doc.processed_at = datetime.utcnow()

        if error:
            doc.error_message = error

        logger.info(f"Updated document {doc.id} status to {status}")
        return doc

    async def upload_source(
        self,
        file: UploadFile,
        category: str | None = None,
        tags: str | None = None,
        owner_id: str | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> SourceDocument:
        """
        Upload and process a source document.

        Args:
            file: Uploaded file
            category: Optional category or topic
            tags: Optional comma-separated tags
            owner_id: Optional owner ID
            metadata: Optional additional metadata

        Returns:
            Created source document

        Raises:
            SourceServiceError: If upload fails

        """
        try:
            # Setup metadata
            processed_tags = []
            if tags:
                processed_tags = [tag.strip() for tag in tags.split(",") if tag.strip()]

            if metadata is None:
                metadata = {}

            # Get content hash for deduplication
            content_hash = await self._compute_content_hash(file.file)

            # Determine MIME type
            content_type = file.content_type or mimetypes.guess_type(file.filename)[0] or "application/octet-stream"

            # Generate source document ID
            source_id = str(uuid.uuid4())

            # Create source document directory
            source_dir = Path(self.upload_dir) / source_id

            # Save file to disk
            try:
                file_path, file_size = await self._save_uploaded_file(file, source_dir, file.filename)
            except Exception as e:
                raise SourceServiceError(f"Failed to save file: {e!s}") from e

            # Determine document type
            try:
                doc_type = self._determine_document_type(file_path, content_type)
            except Exception as e:
                # Clean up on failure
                shutil.rmtree(source_dir, ignore_errors=True)
                raise SourceServiceError(f"Failed to determine document type: {e!s}") from e

            # Create source document
            source_doc = SourceDocument(
                id=source_id,
                filename=file.filename,
                content_hash=content_hash,
                file_type=doc_type,
                mime_type=content_type,
                file_size=file_size,
                file_path=file_path,
                owner_id=owner_id,
                category=category,
                tags=processed_tags,
                metadata=metadata,
                status=DocumentStatus.PENDING,
            )

            logger.info(f"Created source document: {source_id}")

            # Start processing in background - intentionally not awaited
            asyncio.create_task(self._process_document(source_doc))  # noqa: RUF006

            return source_doc

        except SourceServiceError as e:
            logger.error(f"Source upload failed: {e!s}")
            raise
        except Exception as e:
            logger.error(f"Unexpected error in upload_source: {e!s}")
            raise SourceServiceError(f"Unexpected error during upload: {e!s}") from e

    async def _process_document(self, source_doc: SourceDocument) -> None:
        """
        Process a source document in the background.

        Args:
            source_doc: Source document to process

        """
        try:
            # Update status to processing
            source_doc.status = DocumentStatus.PROCESSING
            source_doc.updated_at = datetime.utcnow()

            logger.info(f"Processing document {source_doc.id}: {source_doc.filename}")

            # Extract text
            try:
                text = await self._extract_text_from_file(str(source_doc.file_path), source_doc.file_type)
            except Exception as e:
                await self._update_source_document_status(source_doc, DocumentStatus.FAILED, str(e))
                return

            if not text.strip():
                await self._update_source_document_status(source_doc, DocumentStatus.FAILED, "Extracted text is empty")
                return

            # Prepare metadata for chunks
            chunk_metadata = {
                "source_id": source_doc.id,
                "filename": source_doc.filename,
                "category": source_doc.category,
            }

            if source_doc.metadata:
                # Add document metadata but don't overwrite chunk metadata keys
                for k, v in source_doc.metadata.items():
                    if k not in chunk_metadata:
                        chunk_metadata[k] = v

            # Split into chunks
            try:
                chunks = await self._chunk_text(text, chunk_metadata)
            except Exception as e:
                await self._update_source_document_status(source_doc, DocumentStatus.FAILED, f"Chunking failed: {e!s}")
                return

            # Create embeddings and store in vector database
            try:
                document_chunks, embedding_metadata = await self._store_chunks(source_doc.id, chunks)
            except Exception as e:
                await self._update_source_document_status(
                    source_doc,
                    DocumentStatus.FAILED,
                    f"Vector storage failed: {e!s}",
                )
                return

            # Update document with success status and counts
            source_doc.chunk_count = len(document_chunks)
            source_doc.embedding_count = len(embedding_metadata)
            await self._update_source_document_status(source_doc, DocumentStatus.COMPLETED)

            logger.info(
                f"Document processing completed: {source_doc.id}, "
                f"{source_doc.chunk_count} chunks, "
                f"{source_doc.embedding_count} embeddings"
            )

        except Exception as e:
            logger.error(f"Document processing failed: {e!s}")
            await self._update_source_document_status(source_doc, DocumentStatus.FAILED, str(e))

    async def get_source_document(self, source_id: str) -> SourceDocument | None:
        """
        Get a source document by ID.

        Args:
            source_id: Source document ID

        Returns:
            Source document or None if not found

        """
        # In a real implementation, this would retrieve the document from storage
        # For now, we're returning an example document for illustration
        try:
            # This would typically be a database query
            # For now, we'll just return a mock document
            return SourceDocument(
                id=source_id,
                filename=f"document-{source_id}.pdf",
                content_hash="abc123",
                file_type=DocumentType.PDF,
                mime_type="application/pdf",
                file_size=12345,
                status=DocumentStatus.COMPLETED,
                created_at=datetime.utcnow(),
                updated_at=datetime.utcnow(),
                processed_at=datetime.utcnow(),
                chunk_count=10,
                embedding_count=10,
            )
        except Exception as e:
            logger.error(f"Error getting source document {source_id}: {e!s}")
            return None

    async def list_sources(
        self,
        page: int = 1,  # noqa: ARG002
        page_size: int = 20,  # noqa: ARG002
        status_filter: DocumentStatus | None = None,  # noqa: ARG002
        category: str | None = None,  # noqa: ARG002
        owner_id: str | None = None,  # noqa: ARG002
    ) -> tuple[list[SourceDocumentSummary], int]:
        """
        List source documents with pagination and filtering.

        Args:
            page: Page number (1-based)
            page_size: Items per page
            status_filter: Filter by status
            category: Filter by category
            owner_id: Filter by owner

        Returns:
            Tuple of (list of source documents, total count)

        """
        # In a real implementation, this would query a database
        # For now, we're returning example data

        # Mock data
        docs = [
            SourceDocumentSummary(
                id=f"doc-{i}",
                filename=f"document-{i}.pdf",
                file_type="pdf",
                status=DocumentStatus.COMPLETED,
                created_at=datetime.utcnow(),
                updated_at=datetime.utcnow(),
                chunk_count=10,
                file_size=12345,
            )
            for i in range(1, 6)
        ]

        total = len(docs)
        return docs, total

    async def search_documents(
        self,
        query: str,
        source_ids: list[str] | None = None,
        k: int = 5,
        min_score: float = 0.0,
    ) -> list[dict[str, Any]]:
        """
        Search for documents by semantic similarity.

        Args:
            query: Search query
            source_ids: Optional list of source IDs to search within
            k: Maximum number of results to return
            min_score: Minimum similarity score threshold

        Returns:
            List of search results with scores and metadata

        """
        try:
            # Generate query embedding
            query_embeddings = await self.embedding_service.generate_embeddings([query])

            if not query_embeddings or len(query_embeddings) == 0:
                raise SourceServiceError("Failed to generate query embedding")

            query_embedding = query_embeddings[0]

            # Prepare filter if source_ids provided
            filter_expr = None
            if source_ids:
                filter_expr = {"source_id": {"$in": source_ids}}

            # Search vector database
            await self.vector_store.get_collection(settings.knowledge_config.vector_db.documents_collection)

            results = await self.vector_store.search(
                collection_name=settings.knowledge_config.vector_db.documents_collection,
                query_embedding=query_embedding,
                k=k,
                filter=filter_expr,
                include_metadata=True,
                include_values=False,
            )

            # Format results
            search_results = []
            for result in results:
                if result["score"] < min_score:
                    continue

                search_results.append(
                    {
                        "text": result["text"],
                        "score": result["score"],
                        "metadata": result["metadata"],
                    }
                )

            return search_results

        except Exception as e:
            logger.error(f"Search failed: {e!s}")
            raise SourceServiceError(f"Search failed: {e!s}") from e

    async def delete_source(self, source_id: str) -> bool:
        """
        Delete a source document and all associated data.

        Args:
            source_id: Source document ID

        Returns:
            True if deletion was successful

        Raises:
            SourceServiceError: If deletion fails

        """
        try:
            # Get document to verify it exists
            doc = await self.get_source_document(source_id)
            if not doc:
                raise SourceServiceError(f"Source document not found: {source_id}")

            # Delete from vector database
            try:
                await self.vector_store.get_collection(settings.knowledge_config.vector_db.documents_collection)
                await self.vector_store.delete(
                    collection_name=settings.knowledge_config.vector_db.documents_collection,
                    filter={"source_id": source_id},
                )
            except Exception as e:
                logger.error(f"Vector deletion failed for {source_id}: {e!s}")
                # Continue with other deletion steps

            # Delete file from disk if it exists
            if doc.file_path and Path(doc.file_path).exists():
                try:
                    # Delete the whole document directory
                    source_dir = Path(doc.file_path).parent
                    shutil.rmtree(source_dir)
                except Exception as e:
                    logger.error(f"File deletion failed for {source_id}: {e!s}")

            # Update document status to deleted
            doc.status = DocumentStatus.DELETED
            doc.updated_at = datetime.utcnow()

            logger.info(f"Deleted source document: {source_id}")

            return True

        except SourceServiceError as e:
            logger.error(f"Source deletion failed: {e!s}")
            raise
        except Exception as e:
            logger.error(f"Unexpected error in delete_source: {e!s}")
            raise SourceServiceError(f"Unexpected error during deletion: {e!s}") from e


# Factory function to create source service instance
async def get_source_service():
    """Get a source service instance with dependencies."""
    from app.core.config import settings
    from app.core.text_processing import TextProcessor
    from app.knowledge.infrastructure.vector_store_client import VectorStoreClient
    from app.knowledge.services.embedding_service import EmbeddingService

    # Create dependencies
    text_processor = TextProcessor()
    vector_store = await VectorStoreClient.create()
    embedding_service = EmbeddingService()

    return SourceService(
        text_processor=text_processor,
        embedding_service=embedding_service,
        vector_store_client=vector_store,
        upload_dir=settings.upload_dir,
        openai_api_key=settings.openai_api_key,
    )


# Singleton source service instance
source_service = None


# Get or create the source service singleton
async def get_or_create_source_service():
    """Get or create a source service singleton instance."""
    global source_service  # noqa: PLW0603
    if source_service is None:
        source_service = await get_source_service()
    return source_service

"""Source document service for file upload and processing."""

import asyncio
import contextlib
import hashlib
import logging
import mimetypes
from pathlib import Path
from typing import Any

import aiofiles
import docx
from fastapi import HTTPException, UploadFile
from langchain.text_splitter import RecursiveCharacterTextSplitter
from pypdf import PdfReader

from app.core.vector_config import document_processing_config, vector_db_config
from app.db.vector_store import VectorStoreError, get_vector_store
from app.knowledge.models.chunk import DocumentChunk
from app.knowledge.models.source import DocumentStatus, DocumentType, SourceDocument
from app.knowledge.services.embedding_service import EmbeddingService

logger = logging.getLogger(__name__)


class SourceServiceError(Exception):
    """Base exception for source service operations."""

    pass


class FileProcessingError(SourceServiceError):
    """Raised when file processing fails."""

    pass


class UnsupportedFileTypeError(SourceServiceError):
    """Raised when file type is not supported."""

    pass


class SourceService:
    """Service for managing source documents and file processing."""

    def __init__(self):
        """Initialize the source service."""
        self.upload_dir = Path(document_processing_config.storage_path)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

        self.max_file_size = document_processing_config.max_file_size
        self.allowed_types = document_processing_config.allowed_types
        self.chunk_size = document_processing_config.chunk_size
        self.chunk_overlap = document_processing_config.chunk_overlap

        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""],
        )

        # Initialize embedding service
        self.embedding_service = EmbeddingService()

        logger.info(f"SourceService initialized with upload dir: {self.upload_dir}")

    async def upload_source(
        self,
        file: UploadFile,
        owner_id: str | None = None,
        category: str | None = None,
        tags: list[str] | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> SourceDocument:
        """
        Upload and process a source document.

        Args:
            file: FastAPI UploadFile object
            owner_id: ID of the user uploading the file
            category: Document category
            tags: List of tags for the document
            metadata: Additional metadata

        Returns:
            SourceDocument with initial metadata

        Raises:
            HTTPException: If file validation fails
            SourceServiceError: If processing fails
        """
        try:
            # Validate file
            await self._validate_file(file)

            # Generate file hash and save file
            file_path, content_hash = await self._save_file(file)

            # Create source document record
            source_doc = await self._create_source_record(
                file=file,
                file_path=file_path,
                content_hash=content_hash,
                owner_id=owner_id,
                category=category,
                tags=tags or [],
                metadata=metadata or {},
            )

            logger.info(
                f"File uploaded successfully: {source_doc.filename} ({source_doc.id})"
            )

            # Start background processing
            asyncio.create_task(self._process_source_async(file_path, source_doc))

            return source_doc

        except Exception as e:
            logger.error(f"Failed to upload source file: {str(e)}")

            # Clean up file if it was saved
            if "file_path" in locals():
                with contextlib.suppress(OSError):
                    Path(file_path).unlink()

            if isinstance(e, HTTPException | SourceServiceError):
                raise
            raise SourceServiceError(f"Upload failed: {str(e)}") from e

    async def _validate_file(self, file: UploadFile) -> None:
        """
        Validate uploaded file.

        Args:
            file: UploadFile to validate

        Raises:
            HTTPException: If validation fails
        """
        # Check file size
        if file.size and file.size > self.max_file_size:
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Maximum size: {self.max_file_size / (1024*1024):.1f}MB",
            )

        # Check file extension
        if not file.filename:
            raise HTTPException(status_code=400, detail="Filename is required")

        file_ext = Path(file.filename).suffix.lower().lstrip(".")
        if file_ext not in self.allowed_types:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type. Allowed: {', '.join(self.allowed_types)}",
            )

        # Check MIME type
        mime_type, _ = mimetypes.guess_type(file.filename)
        if mime_type and not self._is_mime_type_allowed(mime_type, file_ext):
            raise HTTPException(
                status_code=400,
                detail=f"MIME type {mime_type} not allowed for .{file_ext} files",
            )

    def _is_mime_type_allowed(self, mime_type: str, file_ext: str) -> bool:
        """Check if MIME type is allowed for the file extension."""
        allowed_mimes = {
            "pdf": ["application/pdf"],
            "txt": ["text/plain", "text/x-plain"],
            "md": ["text/markdown", "text/x-markdown", "text/plain"],
            "docx": [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ],
            "html": ["text/html"],
            "json": ["application/json", "text/json"],
            "csv": ["text/csv", "application/csv"],
        }

        return mime_type in allowed_mimes.get(file_ext, [])

    async def _save_file(self, file: UploadFile) -> tuple[str, str]:
        """
        Save uploaded file and calculate hash.

        Args:
            file: UploadFile to save

        Returns:
            Tuple of (file_path, content_hash)
        """
        # Generate unique filename
        timestamp = int(asyncio.get_event_loop().time())
        safe_filename = f"{timestamp}_{file.filename}"
        file_path = self.upload_dir / safe_filename

        # Calculate hash while saving
        hash_sha256 = hashlib.sha256()

        async with aiofiles.open(file_path, "wb") as f:
            await file.seek(0)  # Reset file pointer
            while chunk := await file.read(8192):  # Read in 8KB chunks
                hash_sha256.update(chunk)
                await f.write(chunk)

        content_hash = hash_sha256.hexdigest()

        logger.debug(f"File saved: {file_path} (hash: {content_hash[:16]}...)")
        return str(file_path), content_hash

    async def _create_source_record(
        self,
        file: UploadFile,
        file_path: str,
        content_hash: str,
        owner_id: str | None = None,
        category: str | None = None,
        tags: list[str] | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> SourceDocument:
        """
        Create initial source document record.

        Args:
            file: Original UploadFile
            file_path: Path where file was saved
            content_hash: SHA-256 hash of file content
            owner_id: Owner user ID
            category: Document category
            tags: Document tags
            metadata: Additional metadata

        Returns:
            Created SourceDocument
        """
        # Get file info
        file_size = file.size or Path(file_path).stat().st_size
        mime_type, _ = mimetypes.guess_type(file.filename)

        # Create document metadata
        doc_metadata = metadata or {}
        doc_metadata.update(
            {
                "original_filename": file.filename,
                "upload_timestamp": asyncio.get_event_loop().time(),
                "mime_type": mime_type,
                "file_size": file_size,
            }
        )

        # Create source document
        return SourceDocument(
            filename=file.filename,
            content_hash=content_hash,
            status=DocumentStatus.PENDING,
            file_size=file_size,
            mime_type=mime_type,
            file_path=file_path,
            collection_name=vector_db_config.documents_collection,
            metadata=doc_metadata,
            tags=tags or [],
            category=category,
            owner_id=owner_id,
        )

        # TODO: Save to database if using ORM
        # For now, we'll just return the in-memory object

    async def _process_source_async(
        self, file_path: str, source_doc: SourceDocument
    ) -> None:
        """
        Asynchronously process source document in background.

        Args:
            file_path: Path to the uploaded file
            source_doc: Source document to process
        """
        try:
            await self._process_source(file_path, source_doc.id)
        except Exception as e:
            logger.error(f"Background processing failed for {source_doc.id}: {str(e)}")
            source_doc.mark_as_failed(str(e))
            # TODO: Update database record

    async def _process_source(self, file_path: str, source_id: str) -> None:
        """
        Process source document: extract text, create chunks, generate embeddings.

        Args:
            file_path: Path to the uploaded file
            source_id: Source document ID
        """
        try:
            # TODO: Get source document from database
            # For now, create a minimal document object
            source_doc = SourceDocument(
                id=source_id,
                filename=Path(file_path).name,
                content_hash="temp_hash",
                file_path=file_path,
            )

            logger.info(f"Starting processing for {source_doc.filename} ({source_id})")

            # Update status to processing
            source_doc.mark_as_processing()
            # TODO: Update database record

            # Extract text content
            content = await self._extract_text_content(
                file_path, source_doc.document_type
            )
            source_doc.content = content
            source_doc.content_length = len(content)

            # Create text chunks
            chunks = await self._create_text_chunks(content, source_id)
            source_doc.chunk_count = len(chunks)

            # Generate embeddings and store in vector database
            embedding_count = await self._process_embeddings(chunks, source_doc)
            source_doc.embedding_count = embedding_count

            # Mark as completed
            source_doc.mark_as_completed(
                chunk_count=len(chunks), embedding_count=embedding_count
            )

            logger.info(
                f"Processing completed for {source_doc.filename}: "
                f"{len(chunks)} chunks, {embedding_count} embeddings"
            )

            # TODO: Update database record

        except Exception as e:
            logger.error(f"Processing failed for {source_id}: {str(e)}")
            # TODO: Update database record to failed status
            raise FileProcessingError(f"Processing failed: {str(e)}") from e

    async def _extract_text_content(
        self, file_path: str, doc_type: DocumentType | None
    ) -> str:
        """
        Extract text content from file based on type.

        Args:
            file_path: Path to the file
            doc_type: Document type

        Returns:
            Extracted text content

        Raises:
            UnsupportedFileTypeError: If file type is not supported
            FileProcessingError: If extraction fails
        """
        try:
            file_ext = Path(file_path).suffix.lower().lstrip(".")

            if file_ext == "pdf" or doc_type == DocumentType.PDF:
                return await self._extract_pdf_content(file_path)
            if (
                file_ext == "txt"
                or doc_type == DocumentType.TXT
                or file_ext in ["md", "markdown"]
                or doc_type == DocumentType.MARKDOWN
            ):
                return await self._extract_text_content_plain(file_path)
            if file_ext == "docx" or doc_type == DocumentType.DOCX:
                return await self._extract_docx_content(file_path)
            if file_ext in ["html", "htm"] or doc_type == DocumentType.HTML:
                return await self._extract_html_content(file_path)
            if file_ext == "json" or doc_type == DocumentType.JSON:
                return await self._extract_json_content(file_path)
            if file_ext == "csv" or doc_type == DocumentType.CSV:
                return await self._extract_csv_content(file_path)
            raise UnsupportedFileTypeError(f"Unsupported file type: {file_ext}")

        except Exception as e:
            if isinstance(e, UnsupportedFileTypeError):
                raise
            raise FileProcessingError(f"Text extraction failed: {str(e)}") from e

    async def _extract_pdf_content(self, file_path: str) -> str:
        """Extract text from PDF file using pypdf."""
        loop = asyncio.get_event_loop()

        def extract_pdf():
            try:
                reader = PdfReader(file_path)
                text_parts = []

                for page_num, page in enumerate(reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_parts.append(f"[Page {page_num + 1}]\n{text}")
                    except Exception as e:
                        logger.warning(
                            f"Failed to extract text from page {page_num + 1}: {str(e)}"
                        )
                        continue

                return "\n\n".join(text_parts)

            except Exception as e:
                raise FileProcessingError(f"PDF processing failed: {str(e)}") from e

        return await loop.run_in_executor(None, extract_pdf)

    async def _extract_text_content_plain(self, file_path: str) -> str:
        """Extract content from plain text file."""
        try:
            async with aiofiles.open(file_path, encoding="utf-8") as f:
                return await f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            async with aiofiles.open(file_path, encoding="latin-1") as f:
                return await f.read()

    async def _extract_docx_content(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        loop = asyncio.get_event_loop()

        def extract_docx():
            try:
                doc = docx.Document(file_path)
                text_parts = []

                for paragraph in doc.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        text_parts.append(text)

                return "\n\n".join(text_parts)

            except Exception as e:
                raise FileProcessingError(f"DOCX processing failed: {str(e)}") from e

        return await loop.run_in_executor(None, extract_docx)

    async def _extract_html_content(self, file_path: str) -> str:
        """Extract text from HTML file."""
        try:
            from bs4 import BeautifulSoup

            async with aiofiles.open(file_path, encoding="utf-8") as f:
                content = await f.read()

            soup = BeautifulSoup(content, "html.parser")

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Extract text
            text = soup.get_text()

            # Clean up whitespace (adapted from WebpageParser.py)
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            return "\n".join(chunk for chunk in chunks if chunk)

        except ImportError as e:
            raise FileProcessingError(
                "BeautifulSoup4 is required for HTML processing"
            ) from e
        except Exception as e:
            raise FileProcessingError(f"HTML processing failed: {str(e)}") from e

    async def _extract_json_content(self, file_path: str) -> str:
        """Extract content from JSON file."""
        import json

        try:
            async with aiofiles.open(file_path, encoding="utf-8") as f:
                content = await f.read()

            data = json.loads(content)

            # Convert JSON to readable text format
            return json.dumps(data, indent=2, ensure_ascii=False)

        except Exception as e:
            raise FileProcessingError(f"JSON processing failed: {str(e)}") from e

    async def _extract_csv_content(self, file_path: str) -> str:
        """Extract content from CSV file."""
        import csv

        try:
            text_parts = []

            async with aiofiles.open(file_path, encoding="utf-8") as f:
                content = await f.read()

            # Parse CSV
            reader = csv.reader(content.splitlines())

            for row_num, row in enumerate(reader):
                if row_num == 0:
                    # Header row
                    text_parts.append(f"Headers: {', '.join(row)}")
                else:
                    # Data rows
                    text_parts.append(f"Row {row_num}: {', '.join(row)}")

            return "\n".join(text_parts)

        except Exception as e:
            raise FileProcessingError(f"CSV processing failed: {str(e)}") from e

    async def _create_text_chunks(
        self, content: str, source_id: str
    ) -> list[DocumentChunk]:
        """
        Split text content into chunks using langchain text splitter.

        Args:
            content: Text content to split
            source_id: Source document ID

        Returns:
            List of DocumentChunk objects
        """
        try:
            # Split text into chunks
            text_chunks = self.text_splitter.split_text(content)

            chunks = []
            current_position = 0

            for chunk_index, chunk_text in enumerate(text_chunks):
                # Find the position of this chunk in the original text
                start_pos = content.find(chunk_text, current_position)
                if start_pos == -1:
                    start_pos = current_position

                end_pos = start_pos + len(chunk_text)
                current_position = end_pos

                # Create chunk object
                chunk = DocumentChunk(
                    source_document_id=source_id,
                    content=chunk_text,
                    chunk_index=chunk_index,
                    start_position=start_pos,
                    end_position=end_pos,
                    content_length=len(chunk_text),
                    metadata={
                        "chunk_size": len(chunk_text),
                        "overlap_with_previous": chunk_index > 0,
                        "text_preview": (
                            chunk_text[:100] + "..."
                            if len(chunk_text) > 100
                            else chunk_text
                        ),
                    },
                )

                chunks.append(chunk)

            logger.info(
                f"Created {len(chunks)} chunks from content ({len(content)} chars)"
            )
            return chunks

        except Exception as e:
            raise FileProcessingError(f"Text chunking failed: {str(e)}") from e

    async def _process_embeddings(
        self, chunks: list[DocumentChunk], source_doc: SourceDocument
    ) -> int:
        """
        Generate embeddings for chunks and store in vector database.

        Args:
            chunks: List of document chunks
            source_doc: Source document

        Returns:
            Number of embeddings created
        """
        try:
            if not chunks:
                return 0

            # Get vector store client
            vector_store = await get_vector_store()

            # Ensure collection exists
            await vector_store.create_collection(
                name=vector_db_config.documents_collection,
                metadata=vector_db_config.get_collection_metadata("documents"),
            )

            # Prepare data for vector store
            documents = [chunk.content for chunk in chunks]
            chunk_ids = [chunk.id for chunk in chunks]
            metadatas = []

            for chunk in chunks:
                metadata = {
                    "source_document_id": source_doc.id,
                    "chunk_id": chunk.id,
                    "chunk_index": chunk.chunk_index,
                    "filename": source_doc.filename,
                    "document_type": (
                        source_doc.document_type.value
                        if source_doc.document_type
                        else "unknown"
                    ),
                    "content_length": chunk.content_length,
                    "start_position": chunk.start_position,
                    "end_position": chunk.end_position,
                    "created_at": chunk.created_at.isoformat(),
                    **chunk.metadata,
                }

                # Add source document metadata
                if source_doc.category:
                    metadata["category"] = source_doc.category
                if source_doc.tags:
                    metadata["tags"] = source_doc.tags
                if source_doc.owner_id:
                    metadata["owner_id"] = source_doc.owner_id

                metadatas.append(metadata)

            # Add documents to vector store (embeddings will be generated automatically by ChromaDB)
            vector_ids = await vector_store.add_documents(
                collection_name=vector_db_config.documents_collection,
                documents=documents,
                metadatas=metadatas,
                ids=chunk_ids,
            )

            # Update chunks with embedding IDs
            for chunk, vector_id in zip(chunks, vector_ids, strict=False):
                chunk.embedding_id = vector_id

            logger.info(
                f"Created {len(vector_ids)} embeddings for {source_doc.filename}"
            )
            return len(vector_ids)

        except VectorStoreError as e:
            raise FileProcessingError(f"Vector store operation failed: {str(e)}") from e
        except Exception as e:
            raise FileProcessingError(f"Embedding processing failed: {str(e)}") from e

    async def search_documents(
        self,
        query: str,
        n_results: int = 5,
        category: str | None = None,
        tags: list[str] | None = None,
        owner_id: str | None = None,
    ) -> dict[str, Any]:
        """
        Search for similar documents using vector search.

        Args:
            query: Search query
            n_results: Number of results to return
            category: Filter by category
            tags: Filter by tags
            owner_id: Filter by owner

        Returns:
            Search results with documents and metadata
        """
        try:
            vector_store = await get_vector_store()

            # Build metadata filter
            where_filter = {}
            if category:
                where_filter["category"] = category
            if owner_id:
                where_filter["owner_id"] = owner_id
            if tags:
                # For tags, we might need a more complex filter
                # This depends on how ChromaDB handles list fields
                where_filter["tags"] = {"$contains": tags[0]} if tags else None

            # Perform vector search
            return await vector_store.search_similar(
                collection_name=vector_db_config.documents_collection,
                query_texts=[query],
                n_results=n_results,
                where=where_filter if where_filter else None,
            )

        except Exception as e:
            logger.error(f"Document search failed: {str(e)}")
            raise SourceServiceError(f"Search failed: {str(e)}") from e


# Global service instance
source_service = SourceService()
